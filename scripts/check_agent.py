"""Check the agent's wiring without needing an API key.

Everything here runs offline except the last check. That is deliberate: the
graph shape, the tool boundaries, the routers and both loop bounds are all
decidable without a model, so they should be checkable without one too.

    python scripts/check_agent.py           # offline only
    python scripts/check_agent.py --live    # plus one real orchestrator call

The --live check is the only one that spends money, and it exercises the one
node that is actually implemented: it asks the orchestrator to route a real
question and prints what it chose and why.
"""

import argparse
import asyncio
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from langchain_core.messages import HumanMessage, ToolMessage
from pydantic import ValidationError

from agent import WORKFLOW, AgentState, NodeSpec, Workflow, build_graph
from agent.nodes import filer as fil
from agent.nodes import orchestrator as orch
from agent.nodes import researcher as res
from agent.nodes import responder as resp
from agent.nodes import verifier as ver
from agent.state import (
    FILER,
    FINISH,
    MAX_DELEGATIONS,
    MAX_RETRIEVALS,
    MAX_REVISIONS,
    ORCHESTRATOR,
    RESPONDER,
    ROUTES,
    WORKERS,
)
from agent.tools import find_section, graph_rag_search, read_section
from agent.tools.tickets import CREATE_TICKET, ticket_tools
from agent.tools.tickets import load_tools as ticket_load_tools

PASSED, FAILED = [], []


def check(label: str, ok: bool, detail: str = "") -> None:
    (PASSED if ok else FAILED).append(label)
    print(f"  [{'ok' if ok else 'FAIL'}] {label}{f' — {detail}' if detail else ''}")


def flat(text: str) -> str:
    """Collapse whitespace, so a phrase check is not defeated by a line wrap."""
    return " ".join(text.split())


def load_cli():
    """Import scripts/agent.py as a module.

    It cannot simply be imported: its file is `agent.py` and it sits next to a
    package also called `agent`, which is already on sys.path and wins. Loading
    it from its path under another name is the only way to check the CLI's own
    rules — and those rules decide what the user is offered, so they are worth
    checking.
    """
    import importlib.util

    path = Path(__file__).with_name("agent.py")
    spec = importlib.util.spec_from_file_location("_cli_under_check", path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def rejects(thunk) -> tuple[bool, str]:
    """True if the call raises a validation error, plus the message."""
    try:
        thunk()
    except (ValidationError, ValueError) as error:
        message = error.errors()[0]["msg"] if isinstance(error, ValidationError) else str(error)
        return True, message
    return False, "accepted, should not have"


# ---------------------------------------------------------------- structure


def check_structure() -> None:
    print("\nGraph structure")
    graph = build_graph().get_graph()
    names = set(graph.nodes)
    expected = {spec.name for spec in WORKFLOW.nodes}
    check("every declared node is in the compiled graph", expected <= names)
    check(
        "nodes with tools got their own executor",
        all(
            spec.tools_node in names
            for spec in WORKFLOW.nodes
            if spec.tools
        ),
    )

    edges = {(edge.source, edge.target) for edge in graph.edges}
    check("START enters at the orchestrator", ("__start__", ORCHESTRATOR) in edges)
    check(
        "orchestrator can reach every worker",
        all((ORCHESTRATOR, worker) in edges for worker in WORKERS),
    )
    check(
        "orchestrator cannot reach the verifier directly",
        (ORCHESTRATOR, "verifier") not in edges,
        "the answer is only ever graded after the responder writes it",
    )
    check("responder always hands to the verifier", (RESPONDER, "verifier") in edges)
    check("the verifier reports back to the orchestrator", ("verifier", ORCHESTRATOR) in edges)
    check(
        "the verifier cannot end the run",
        ("verifier", "__end__") not in edges,
        "it grades; the orchestrator decides when grading is over",
    )
    check(
        "the orchestrator can end the run",
        (ORCHESTRATOR, "__end__") in edges,
        f"via the {FINISH!r} route",
    )
    check(
        "the filer ends the run itself",
        (FILER, "__end__") in edges,
        "it acts rather than gathers, so there is nothing to hand back",
    )
    check(
        "the orchestrator and the filer are the only ways out",
        {source for source, target in edges if target == "__end__"}
        == {ORCHESTRATOR, FILER},
        "one exit for answers, one for actions, and no third opinion",
    )
    check(
        "the filer never routes onward to the responder or the verifier",
        not any((FILER, target) in edges for target in (RESPONDER, "verifier")),
        "a ticket confirmation is not an answer to be graded for citations",
    )


def check_tool_boundaries() -> None:
    print("\nPer-node tool access")
    by_name = WORKFLOW.by_name
    check(
        "the researcher has find, read and the graph",
        [tool.name for tool in by_name["researcher"].tools]
        == ["find_section", "read_section", "graph_rag_search"],
    )
    check(
        "the filer has the MCP ticket tool",
        CREATE_TICKET in [tool.name for tool in by_name[FILER].tools],
        f"discovered: {[tool.name for tool in by_name[FILER].tools]}",
    )
    check(
        "the filer cannot retrieve",
        not ({"graph_rag_search", "find_section", "read_section"}
             & {tool.name for tool in by_name[FILER].tools}),
        "it fills the ticket from the conversation, not from a fresh search",
    )
    check(
        "no node but the filer can write a ticket",
        [spec.name for spec in WORKFLOW.nodes
         if CREATE_TICKET in {tool.name for tool in spec.tools}] == [FILER],
        "the one tool with a side effect is listed exactly once",
    )
    check(
        "orchestrator, responder and verifier have no tools",
        not any(by_name[name].tools for name in (ORCHESTRATOR, RESPONDER, "verifier")),
    )


# ------------------------------------------------------------------- routing


def check_routers() -> None:
    print("\nRouters")
    from agent import graph as wiring
    from agent.graph import ROUTE_TARGETS

    check(
        "there is exactly one router",
        [name for name in vars(wiring) if name.startswith("_route")]
        == ["_route_to_worker"],
        "a second node deciding where control goes is a second place the loop "
        "bounds have to be re-derived, and last time they were not",
    )
    check(
        "no choice falls through to the responder",
        orch_route(AgentState()) == RESPONDER,
    )
    check(
        "a choice is honoured",
        orch_route(AgentState(next_node="researcher")) == "researcher",
    )
    check(
        "every route has a target",
        set(ROUTE_TARGETS) == set(ROUTES),
        f"routes: {sorted(ROUTES)}",
    )
    check(
        "finish is the route that ends the run",
        ROUTE_TARGETS[FINISH] == "__end__",
    )
    check(
        "every worker routes to itself",
        all(ROUTE_TARGETS[worker] == worker for worker in WORKERS),
    )


def orch_route(state: AgentState) -> str:
    from agent.graph import _route_to_worker

    return _route_to_worker(state)


def check_orchestrator() -> None:
    print("\nOrchestrator")
    check(
        "its prompt describes every worker it can pick",
        all(worker in orch.PROMPT for worker in WORKERS),
    )
    check(
        "its decision is constrained to an enum",
        orch.Route.model_json_schema()["properties"]["next_node"]["enum"] == list(WORKERS),
    )
    check(
        "the model cannot choose to stop",
        FINISH not in orch.Route.model_json_schema()["properties"]["next_node"]["enum"],
        "stopping is read off the verdict and the counters, not asked for",
    )
    check(
        "it works out the task before picking who does it",
        list(orch.Route.model_fields) == ["brief", "next_node"],
        "field order is what stops the choice being a first impression",
    )
    check(
        "a delegation carries a brief to the worker",
        "[orchestrator → researcher] find the exact title"
        in orch._delegation("researcher", "find the exact title").content,
        "a worker re-run with no task does nothing, seven times over",
    )
    check(
        "the prompt demands a NEW target on a repeat delegation",
        "must name what is new" in flat(orch.PROMPT),
    )

    capped = AgentState(delegations=["researcher"] * MAX_DELEGATIONS)
    result = asyncio.run(orch.run(capped))
    check(
        "the delegation cap short-circuits with no model call",
        result["next_node"] == RESPONDER,
        f"MAX_DELEGATIONS={MAX_DELEGATIONS}",
    )
    check(
        "even the capped delegation tells the responder why",
        "budget" in result["messages"][0].content,
    )
    check(
        "it is told what it has already run",
        "researcher" in orch._progress(AgentState(delegations=["researcher"])),
    )
    check(
        "it is told when the verifier rejected the answer",
        "rejected" in orch._progress(AgentState(verdict="fail", revisions=1)),
    )
    check_finalizer()


def check_finalizer() -> None:
    """The termination rules, which moved here from the verifier's own router.

    All of it is decidable without a model, and all of it runs before any model
    call in orch.run — so these assertions exercise the real code path rather
    than an offline approximation of it.
    """
    print("\nOrchestrator as finalizer")
    graded = dict(delegations=["researcher", RESPONDER], answer="An answer.")

    check(
        "an ungraded run is not finished",
        not orch._is_done(AgentState(**graded)),
        "verdict is None for the whole gathering phase",
    )
    check(
        "a passing answer finishes the run",
        orch._is_done(AgentState(**graded, verdict="pass")),
    )
    check(
        "an answer the grader could not grade finishes the run",
        orch._is_done(AgentState(**graded, verdict="ungraded")),
        "nothing about it is known to be wrong, so a revision would buy nothing",
    )
    check(
        "a failing answer goes back for another pass",
        not orch._is_done(AgentState(**graded, verdict="fail", revisions=0)),
    )
    check(
        "a failing answer stops at the revision cap",
        orch._is_done(AgentState(**graded, verdict="fail", revisions=MAX_REVISIONS)),
        f"MAX_REVISIONS={MAX_REVISIONS}",
    )

    passed = AgentState(**graded, verdict="pass", ticket_recipient="Duty Manager")
    result = asyncio.run(orch.run(passed))
    check(
        "finishing costs no model call",
        result["next_node"] == FINISH,
        "the terminal branch returns before the router is built",
    )
    check(
        "it composes the reply the caller will show",
        "Duty Manager" in result["reply"] and result["reply"].startswith("An answer."),
        "the answer, then the ticket offer",
    )
    check(
        "it does not touch the graded answer",
        "answer" not in result,
        "the verifier graded that exact string; appending an offer to it would "
        "mean the text that was checked is not the text that goes out",
    )
    check(
        "finishing delegates to nobody and says nothing",
        "delegations" not in result and "messages" not in result,
    )

    # The ordering bug this design has to avoid: now that the verifier hands
    # back, a run that is both graded and out of delegations arrives here again.
    # If the cap were read first it would answer, be graded, and be sent to the
    # responder once more, forever.
    capped_and_passed = AgentState(
        delegations=["researcher"] * MAX_DELEGATIONS,
        answer="An answer.",
        verdict="pass",
    )
    result = asyncio.run(orch.run(capped_and_passed))
    check(
        "a finished run at the delegation cap still finishes",
        result["next_node"] == FINISH,
        "the terminal branch is checked before the cap, and that order is the bug",
    )


def check_verifier() -> None:
    print("\nVerifier")
    check(
        "its verdict is constrained to pass or fail",
        ver.Verdict.model_json_schema()["properties"]["verdict"]["enum"] == ["pass", "fail"],
    )
    check(
        "it has to enumerate problems before ruling",
        list(ver.Verdict.model_fields)[0] == "unsupported_claims",
        "field order is what stops a verdict-first first impression",
    )
    check(
        "an honest 'not in the documents' answer is a documented pass",
        "documents do not cover the question" in flat(ver.PROMPT),
        "without this the loop never terminates on an unanswerable question",
    )

    # Reachable without a model: no answer to grade short-circuits.
    result = asyncio.run(ver.run(AgentState()))
    check("a missing answer fails without crashing", result["verdict"] == "fail")
    check("a missing answer still counts as a revision", result["revisions"] == 1)

    critique = ver._critique(
        ver.Verdict(
            unsupported_claims=["the title 'Assistant Manager'"],
            verdict="fail",
            gap="confirm the actual title that approves refunds",
        )
    )
    check("a rejection names the claim", "Assistant Manager" in critique)
    check("a rejection says what would close it", "confirm the actual title" in critique)

    # The mechanical citation check — decidable by string comparison, so it runs
    # before any model call and cannot be talked round.
    evidence = ToolMessage(
        tool_call_id="1",
        content=(
            "[13] sample_role_duties.pdf › page 2 › Shift Supervisor › Key Responsibilities\n"
            "[16] sample_role_duties.pdf › page 1 › Store Associate - Front of House › Out of Scope\n"
        ),
    )
    real = "sample_role_duties.pdf › page 2 › Shift Supervisor › Key Responsibilities"
    shifted = "sample_role_duties.pdf › page 3 › Shift Supervisor › Key Responsibilities"
    check(
        "a copied label passes",
        not ver._citation_problems(f"x [1]\n\n### References\n- [1] {real}\n", [evidence]),
    )
    check(
        "a label with a shifted page is caught",
        shifted
        in str(ver._citation_problems(f"x [1]\n\n### References\n- [1] {shifted}\n", [evidence])),
        "the exact failure two rounds of prompting did not stop",
    )
    stripped = "Shift Supervisor › Key Responsibilities"
    check(
        "a label stripped of its file and page is caught",
        stripped
        in str(ver._citation_problems(f"x [1]\n\n### References\n- [1] {stripped}\n", [evidence])),
        "the loophole the responder found when told not to fabricate",
    )
    check(
        "formatting alone does not fail a correct citation",
        not ver._citation_problems(f"x [1]\n\n### References\n- [1] **{real}**.\n", [evidence]),
        "bold and a full stop are decoration, not a different label",
    )
    check(
        "a trailing colon does not fail a correct citation",
        not ver._citation_problems(
            f"x [1]\n\n### References\n- [1] {real}:\n", [evidence]
        ),
        "several real section labels end in one and answers drop it",
    )
    # Reference numbering style must not decide whether the guard runs at all.
    for style in (f"- [1] {shifted}", f"[1] {shifted}", f"1. {shifted}", f"1) {shifted}"):
        check(
            f"the guard runs on references written as {style.split(' ')[0]!r}",
            shifted in str(ver._citation_problems(f"x\n\n### References\n{style}\n", [evidence])),
            "matching one style only makes the check stop silently, not get stricter",
        )
    check(
        "with no evidence at all, nothing is called fabricated",
        not ver._citation_problems(f"x\n\n### References\n- [1] {real}\n", []),
        "otherwise an honest 'found nothing' answer could never pass",
    )
    # The class of bug, not just the instance: a References section the parser
    # cannot read must fail, because "found no citations" and "found them all
    # correct" are otherwise the same empty result.
    check(
        "an unreadable References section fails instead of passing",
        ver._citation_problems(
            "x\n\n### References\n| 1 | some label |\n", [evidence]
        ),
        "an unrecognised format must not silently switch the check off",
    )
    check(
        "an empty References section fails too",
        ver._citation_problems("x\n\n### References\n", [evidence]),
    )
    check(
        "an answer with no References section is left to the grader",
        not ver._citation_problems("x with no references at all", [evidence]),
        "the prompt already covers a missing section; this check is about unreadable ones",
    )
    check(
        "a rejection hands over the labels that ARE valid",
        real
        in asyncio.run(
            ver.run(
                AgentState(
                    answer=f"x [1]\n\n### References\n- [1] {stripped}\n",
                    messages=[evidence],
                )
            )
        )["messages"][0].content,
        "so the fix is copying from a list, not guessing again",
    )
    fabricated_run = asyncio.run(
        ver.run(
            AgentState(
                answer=f"x [1]\n\n### References\n- [1] {shifted}\n", messages=[evidence]
            )
        )
    )
    check(
        "a fabricated citation fails without a model call",
        fabricated_run["verdict"] == "fail",
    )
    check(
        "and the rejection quotes the offending label",
        shifted in fabricated_run["messages"][0].content,
    )


def check_responder() -> None:
    print("\nResponder")
    check("it has no retrieval tools", not WORKFLOW.by_name[RESPONDER].tools)
    check(
        "it is told to repair rather than reword on a revision",
        "Revision policy" in flat(resp.PROMPT),
    )
    check("it is required to cite", "### References" in flat(resp.PROMPT))
    check(
        "'the documents do not cover this' is a complete answer",
        "complete and correct answer" in flat(resp.PROMPT),
        "the responder and verifier have to agree on this or the loop spins",
    )
    check(
        "it is told it has no retrieval of its own",
        "no retrieval tools" in flat(resp.PROMPT),
    )
    check(
        "it is told to copy labels, never build them",
        "Never construct one" in flat(resp.PROMPT),
        "a plausible label points nowhere and cannot be spotted",
    )
    check(
        "the label example is schematic, not a real filename",
        "<file>.pdf" in flat(resp.PROMPT)
        and "sample_role_duties.pdf › page 2 ›" not in flat(resp.PROMPT),
        "a realistic example gets pattern-filled instead of copied",
    )
    check(
        "it is told a bare heading is not a label",
        "is not a label" in flat(resp.PROMPT),
    )
    check(
        "it is warned the corpus holds several organisations",
        "SEVERAL UNRELATED ORGANISATIONS" in flat(resp.PROMPT)
        and "SEVERAL UNRELATED ORGANISATIONS" in flat(res.PROMPT),
        "one graph holds 8 orgs, so unscoped retrieval crosses them",
    )
    check(
        "the verifier checks the same thing the responder is asked for",
        "full source label" in flat(ver.PROMPT) and "full source label" in flat(resp.PROMPT),
    )


# ---------------------------------------------------------------- validation


def check_declarations() -> None:
    print("\nDeclarations that should be refused")
    for label, thunk in [
        (
            "a node name that is not an identifier",
            lambda: NodeSpec(name="gap auditor", system_prompt="x", runner=orch.run),
        ),
        (
            "a reserved node name",
            lambda: NodeSpec(name="__start__", system_prompt="x", runner=orch.run),
        ),
        (
            "the same tool listed twice",
            lambda: NodeSpec(
                name="x", system_prompt="x", runner=orch.run,
                tools=[find_section, find_section],
            ),
        ),
        (
            "a runner that is not callable",
            lambda: NodeSpec(name="x", system_prompt="x", runner="nope"),
        ),
        (
            "a workflow missing a node the wiring needs",
            lambda: Workflow(nodes=[orch.SPEC]),
        ),
        (
            "an unknown state field",
            lambda: AgentState(mesages=[]),
        ),
    ]:
        refused, message = rejects(thunk)
        check(label, refused, message)


def check_tool_arguments() -> None:
    print("\nTool arguments that should be refused")
    for label, tool, args in [
        ("a graph mode that would smuggle in vector search", graph_rag_search, {"question": "q", "mode": "mix"}),
        ("top_k below the floor", find_section, {"question": "q", "top_k": 0}),
        ("top_k above the ceiling", find_section, {"question": "q", "top_k": 99}),
        ("an empty question", find_section, {"question": "", "top_k": 3}),
        ("an argument the tool does not have", find_section, {"question": "q", "limit": 3}),
        ("a read with no id at all", read_section, {"section_id": ""}),
    ]:
        refused, _ = rejects(lambda: asyncio.run(tool.ainvoke(args)))
        check(label, refused)


# -------------------------------------------------------------- behavioural


def check_retrieval_cascade() -> None:
    """Find, read, then escalate — and the three places that must agree on it.

    The order the tools are declared in (which is the order the model sees
    them), the policy in the researcher's prompt, and the tools' own
    descriptions. A tool that describes itself as the strongest option will be
    reached for first no matter what the prompt above it says.
    """
    print("\nRetrieval cascade (find, read, graph on escalation)")
    names = [tool.name for tool in WORKFLOW.by_name["researcher"].tools]
    check(
        "the shortlist tool is declared first",
        names[0] == "find_section",
        f"declared order: {names}",
    )
    check(
        "reading comes before the graph",
        names.index("read_section") < names.index("graph_rag_search"),
    )
    check(
        "exact-term matching is not a tool that can be skipped",
        "keyword_search" not in names,
        "it runs inside find_section on every call — see doctree/search.py",
    )
    check(
        "the prompt opens on the shortlist tool",
        "ALWAYS start with find_section" in flat(res.PROMPT),
    )
    check(
        "the prompt says a preview is not evidence",
        "The preview is NOT evidence" in flat(res.PROMPT)
        and "citing a preview" in flat(res.PROMPT),
        "the one failure this split introduces: answering off the shortlist",
    )
    check(
        "the prompt says to read the role, not just the sub-section",
        "Read the ROLE, not just the sub-section" in flat(res.PROMPT),
        "exclusions mean something different next to the duties they carve out",
    )
    check(
        "the graph is named as the escalation, not the opening move",
        "escalate to graph_rag_search" in flat(res.PROMPT),
    )
    check(
        "confirming an answer it already has is called out as the expensive mistake",
        "Do not escalate out of thoroughness" in flat(res.PROMPT),
        "an unnecessary graph call is several model calls",
    )

    # The tool docstrings are what the model receives as the tool description,
    # so they are part of the policy whether or not they are meant to be.
    check(
        "the shortlist tool describes itself as first",
        "START HERE" in flat(find_section.description),
    )
    check(
        "the shortlist tool refuses to be treated as evidence",
        "Do not answer or cite from the preview" in flat(find_section.description),
    )
    check(
        "the read tool describes itself as the evidence step",
        "The evidence step" in flat(read_section.description),
    )
    check(
        "the graph tool describes itself as escalation",
        "ESCALATION tool, not the opening move" in flat(graph_rag_search.description),
    )
    check(
        "the graph tool sends failures back to the section search, not to itself",
        "go back to find_section" in flat(graph_rag_search.description),
        "retrying the graph with the same wording is what cannot help",
    )
    check(
        "an id is not offered as a citation",
        "Not a citation" in read_section.args_schema.model_fields["section_id"].description
        and "Report the LABEL" in flat(res.PROMPT),
        "the id names a section; the label cites one",
    )

    # Most labels now come from the section store, so what it writes at ingest is
    # what most citations will be. Both writers must accept a label shorter than
    # the canonical shape without padding it out — invented pages are the failure
    # this project has already paid for twice.
    for label, prompt in (("researcher", res.PROMPT), ("responder", resp.PROMPT)):
        check(
            f"the {label} is told a short label is not an incomplete one",
            "short label is not an incomplete one" in flat(prompt)
            or "copy those as they are" in flat(prompt),
        )
    check(
        "the responder is told never to extend one",
        "Never extend it" in flat(resp.PROMPT),
        "adding a plausible page is the fabrication this guards",
    )

    # The label a tool prints has to survive being harvested back out of the tool
    # output, because that round trip is the whole citation check. Anything
    # printed on the same line AFTER a label is absorbed into it — which is why
    # find_section puts the label alone on its line and the id underneath.
    real = "x.pdf › page 3 › Housekeeping Supervisor › Out of Scope:"
    block = (
        f"[1] {real}\n"
        "    id: x#housekeeping-supervisor/out-of-scope\n"
        "    found by: headings+keyword | 48 words\n"
        "    preview: some section text"
    )
    harvested = ver._known_labels([ToolMessage(block, tool_call_id="1")])
    check(
        "a harvested section label carries nothing but the label",
        harvested == {real},
        f"harvested: {sorted(harvested)}",
    )


def check_researcher() -> None:
    print("\nResearcher")
    check("it has all three retrievers", len(WORKFLOW.by_name["researcher"].tools) == 3)
    check(
        "it is told not to answer",
        "You do not answer the question" in flat(res.PROMPT),
        "answering here would leave the verifier grading a summary",
    )
    check(
        "it must confirm graph titles against document text",
        "confirm that exact string with find_section" in flat(res.PROMPT),
    )
    check("its report has to name gaps", "Gaps —" in flat(res.PROMPT))
    check(
        "it is told the brief overrides its own sense of being finished",
        "overrides your own sense of whether the work is finished" in flat(res.PROMPT),
        "this is what stopped it no-opping on re-delegation",
    )
    check(
        "it is told where the graph tool hides its labels",
        "Reference Document List at the END" in flat(res.PROMPT),
        "the chunks carry a reference_id, not a label",
    )
    check(
        "it is warned the bracket means something different per tool",
        "is NOT a citation number" in flat(res.PROMPT),
    )

    check(
        "the tool budget is stated every turn",
        "retrievals left" in res._budget(3),
    )
    check(
        "a spent budget is stated as spent",
        "budget is spent" in res._budget(0),
    )
    check(
        "past the budget the tools are unbound, not just discouraged",
        res._get_llm(tools_allowed=False) is not res._get_llm(tools_allowed=True),
        f"MAX_RETRIEVALS={MAX_RETRIEVALS}",
    )
    check(
        "completed tool calls are what count against the budget",
        res._retrievals_used(
            [HumanMessage("q"), ToolMessage("out", tool_call_id="1"), HumanMessage("x")]
        )
        == 1,
    )


def check_corpora() -> None:
    """The two retrieval stores must hold the same documents.

    Nothing in the code enforces this — they are ingested by separate scripts
    into separate directories — so it is checked rather than assumed. A
    mismatch is not a subtle degradation: the researcher would cross-check the
    graph against a passage from a document the graph has never seen, and the
    responder would cite both in one answer as though they described one
    organisation.
    """
    print("\nCorpus agreement between the two stores")
    try:
        import graph_rag.extraction as graph_extraction
        from doctree import corpus
        from doctree.tree import Node
        from graph_rag.graph_rag import WORKING_DIR

        store = corpus()
        sections = {node.source for node in store.nodes.values()}
        chunks = json.loads(
            (Path(WORKING_DIR) / "kv_store_text_chunks.json").read_text(encoding="utf-8")
        )
        labels = [value.get("file_path", "") for value in chunks.values()]
    except Exception as error:  # noqa: BLE001 - reported, not swallowed
        check("both stores are readable", False, f"{type(error).__name__}: {error}")
        return

    # A label starts with the filename it came from. Ones that do not are not
    # missing documents — they are chunks whose label the extractor built wrong,
    # which is a separate problem and gets its own check below.
    heads = [label.split("›")[0].strip() for label in labels]
    graph = {head for head in heads if head.endswith(".pdf")}
    malformed = sorted({label for label, head in zip(labels, heads) if not head.endswith(".pdf")})

    check(
        "both stores are readable",
        True,
        f"trees={len(sections)} graph={len(graph)} files, {len(store)} sections",
    )
    only_graph = sorted(graph - sections)
    only_trees = sorted(sections - graph)
    check(
        "the section store and the graph store hold the same documents",
        not only_graph and not only_trees,
        f"graph-only: {only_graph or 'none'} | trees-only: {only_trees or 'none'}",
    )
    if only_graph or only_trees:
        print("      fix: python scripts/ingest.py   (re-parses data/raw into data/tree)")

    # A citation label with no file or page is one the responder cannot cite and
    # the verifier cannot check. The chunk is still retrievable, so this degrades
    # an answer rather than breaking a run — but silently.
    check(
        "every graph chunk has a citable label",
        not malformed,
        f"{len(malformed)} without a file prefix: {malformed[:2] or 'none'}",
    )

    # The same demand of the section store. A label is built from the node's own
    # source and heading path, so this cannot fail the way the old one could —
    # but it can still fail on a node the parse left with no page and no
    # headings, which is a section nothing can cite even though it is retrievable.
    uncitable = [
        node
        for node in store.nodes.values()
        if node.text and not node.label.startswith(node.source)
    ]
    check(
        "every indexed section has a citable label",
        not uncitable,
        f"{len(uncitable)} of {len(store)} uncitable"
        + (" — re-run: python scripts/ingest.py" if uncitable else ""),
    )
    check(
        "both stores label sections with the same function",
        Node.label.fget.__globals__["section_label"] is graph_extraction.section_label,
        "separate implementations are what let the two ingests drift apart",
    )
    # An id has to resolve, or a search hands the model something read_section
    # cannot open. The index and the tree are written by one call, so this
    # catches the case where they were written at different times.
    try:
        from agent.tools._stores import get_heading_collection

        indexed = set(get_heading_collection().get()["ids"])
    except Exception as error:  # noqa: BLE001 - reported, not swallowed
        check("the heading index is readable", False, f"{type(error).__name__}: {error}")
        return
    dangling = sorted(indexed - set(store.nodes))
    check(
        "every indexed id resolves to a section",
        not dangling,
        f"{len(dangling)} dangling: {dangling[:2] or 'none'}"
        + (" — re-run: python scripts/ingest.py" if dangling else ""),
    )


def check_retrieval() -> None:
    """Search and read, without an API key.

    `find_sections` is called directly rather than through the tool, with the
    heading half left out. That is the point of it taking `headings` as an
    argument: the embedding call is the only part that needs a key, so leaving
    it out exercises the BM25 half, the fusion and the id round trip offline —
    which is most of what can break here.
    """
    print("\nRetrieval (no key: BM25 half of the section search, and a read)")
    try:
        from doctree import corpus, find_sections

        store = corpus()
        hits = find_sections("refund approval threshold", top_k=2)
    except Exception as error:  # noqa: BLE001 - reported, not swallowed
        check("the section search runs", False, f"{type(error).__name__}: {error}")
        return

    check("the section search returns hits", bool(hits), f"{len(store)} sections indexed")
    if not hits:
        print("      fix: python scripts/ingest.py")
        return

    check(
        "every hit carries an id that resolves",
        all(store.get(hit.id) is not None for hit in hits),
    )
    check(
        "every hit carries a label starting with its file",
        all(hit.label.startswith(hit.node.source) for hit in hits),
        hits[0].label,
    )
    check(
        "a hit says which retriever found it",
        all(hit.found_by for hit in hits),
        f"found by: {hits[0].found_by}",
    )

    # The read is the other half of the split, and the property that matters is
    # that it returns MORE than the search did — a section whole, not a preview.
    whole = store.read(hits[0].id)
    check(
        "reading a section returns the whole thing",
        len(whole) >= len(hits[0].node.text),
        f"{len(hits[0].node.text)} chars of own text -> {len(whole)} with sub-sections",
    )
    check(
        "an unknown id is refused rather than guessed at",
        "[no section with id" in asyncio.run(
            read_section.ainvoke({"section_id": "no-such-document#nowhere"})
        ),
    )
    for hit in hits:
        print(f"      [{'+'.join(hit.found_by)}] {hit.label}")
        print(f"          {hit.id}")


def check_full_lap() -> None:
    """Drive the whole topology with fake bodies, so no model is involved.

    This is the check that proves the edges work: a delegation, a hand-back, an
    answer, a rejection, the verdict returning to the orchestrator, and the
    revision cap ending it there.

    Only the model calls are faked. The termination rule is the real one — the
    fake orchestrator calls `orch._is_done` and `orch._finish` rather than
    reimplementing them, because a lap that terminates by its own private rule
    proves the edges and nothing about the rule the agent actually runs.
    """
    print("\nA full lap with substituted node bodies")
    trace = []

    def recorder(name: str, update: dict):
        async def run(state: AgentState) -> dict:
            trace.append(name)
            return update
        return run

    async def orchestrator(state: AgentState) -> dict:
        """Delegate to the researcher once, then write the answer, then stop.

        The first pass has to be a real delegation, otherwise the run goes
        straight to the responder and the worker hand-back edge — the one thing
        a lap is meant to prove — is never crossed.
        """
        trace.append(ORCHESTRATOR)
        if orch._is_done(state):
            trace.append(FINISH)
            return orch._finish(state)
        first_pass = "researcher" not in state.delegations
        target = "researcher" if first_pass else RESPONDER
        return {"next_node": target, "delegations": [target]}

    async def responder(state: AgentState) -> dict:
        trace.append(RESPONDER)
        return {"answer": f"draft {trace.count(RESPONDER)}"}

    async def verifier(state: AgentState) -> dict:
        trace.append("verifier")
        # Reads state.answer, like the real one — so a responder that wrote
        # nowhere the verifier looks would show up here.
        assert state.answer, "the verifier saw no answer"
        return {"verdict": "fail", "revisions": state.revisions + 1}

    swapped = {
        ORCHESTRATOR: orchestrator,
        "researcher": recorder("researcher", {}),
        RESPONDER: responder,
        "verifier": verifier,
    }

    workflow = Workflow(
        nodes=[
            NodeSpec(
                name=spec.name,
                system_prompt=spec.system_prompt,
                runner=swapped.get(spec.name, recorder(spec.name, {})),
                tools=[],  # no tools: nothing here should reach a real retriever
            )
            for spec in WORKFLOW.nodes
        ]
    )
    result = asyncio.run(workflow.compile().ainvoke(AgentState()))

    print(f"      path: {' -> '.join(trace)}")
    check("the run terminated", True)
    check(
        "a delegated worker handed control back",
        trace[:3] == [ORCHESTRATOR, "researcher", ORCHESTRATOR],
    )
    check(
        "the verify loop is bounded",
        trace.count("verifier") == MAX_REVISIONS,
        f"verifier ran {trace.count('verifier')}x, cap is {MAX_REVISIONS}",
    )
    check("the answer was always graded", trace.count(RESPONDER) == trace.count("verifier"))
    check("revisions were counted", result["revisions"] == MAX_REVISIONS)
    check(
        "a rejected answer is still returned",
        AgentState(**result).answer == f"draft {MAX_REVISIONS}",
        "the cap ends the loop, so the caller gets the last attempt",
    )
    check(
        "every verdict went back to the orchestrator",
        all(
            trace[position + 1] == ORCHESTRATOR
            for position, name in enumerate(trace)
            if name == "verifier"
        ),
        "the verifier reports; it no longer has an exit of its own",
    )
    check(
        "the run ended at the orchestrator's finish route",
        trace[-2:] == [ORCHESTRATOR, FINISH],
    )
    check(
        "the caller was left a reply to show",
        AgentState(**result).reply == f"draft {MAX_REVISIONS}",
        "composed on the way out, with no offer on a rejected answer",
    )


# ------------------------------------------------------------------- tickets


def check_tickets() -> None:
    """The MCP server, the document it writes, and the node that calls it.

    All of it offline: the server is a local subprocess and the model is never
    involved. What this proves is that the one tool with a side effect really
    starts, really writes the four fields a ticket has to carry, and really
    refuses the calls that would produce a ticket nobody can act on.
    """
    print("\nTickets (MCP server + Word document)")
    import tempfile

    import docx

    from ticket_mcp import Ticket, write_ticket

    # --- the document -----------------------------------------------------
    with tempfile.TemporaryDirectory() as workspace:
        directory = Path(workspace)
        ticket = Ticket(
            addressed_to="Housekeeping Supervisor (Housekeeping)",
            subject="Knife found in guest room",
            what_happened="A housekeeper found a knife while servicing a room.",
            situation_summary="The item may have to be addressed with the Police.",
            next_steps=("Leave the item in place.", "Report it to the Supervisor."),
            raised_by="Housekeeper",
            organisation="Hotel",
            priority="urgent",
            references=("sample_role_duties_hotel.pdf › page 3 › Housekeeping Supervisor",),
        )
        path = write_ticket(ticket, directory)
        check("a ticket is written as a .docx", path.suffix == ".docx" and path.exists())

        document = docx.Document(path)
        text = "\n".join(p.text for p in document.paragraphs)
        cells = [cell.text for row in document.tables[0].rows for cell in row.cells]
        check(
            "it names who it is addressed to",
            ticket.addressed_to in cells,
            ticket.addressed_to,
        )
        for label, value in [
            ("what happened", ticket.what_happened),
            ("a summary of the situation", ticket.situation_summary),
            ("the next steps", ticket.next_steps[0]),
            ("the source labels behind it", ticket.references[0]),
        ]:
            check(f"it carries {label}", value in text)

        # A second ticket in the same second must not silently replace the
        # first: the id is only precise to the second, and the one that would
        # disappear is an earlier report of the same incident.
        twin = write_ticket(ticket, directory)
        check(
            "two tickets filed in the same second both survive",
            twin != path and twin.exists() and path.exists(),
            twin.name,
        )

    required = dict(
        addressed_to="X", subject="Y", what_happened="Z", situation_summary="W"
    )
    for label, bad in [
        ("a ticket with no recipient", {"addressed_to": "  "}),
        ("a ticket with no account of what happened", {"what_happened": ""}),
        ("an invented priority", {"priority": "whenever"}),
    ]:
        refused, message = rejects(lambda bad=bad: Ticket(**(required | bad)))
        check(f"{label} is refused", refused, message)

    # --- the MCP server ---------------------------------------------------
    tools = ticket_tools()
    check(
        "the ticket server starts and offers its tool over MCP",
        CREATE_TICKET in [tool.name for tool in tools],
        f"offers {[tool.name for tool in tools]}",
    )

    tool = next(tool for tool in tools if tool.name == CREATE_TICKET)
    schema = tool.args_schema or {}
    properties = schema.get("properties", {})
    check(
        "its schema came from the server, not from a copy in this repo",
        set(properties) >= {"addressed_to", "what_happened", "situation_summary", "next_steps"},
        f"fields: {sorted(properties)}",
    )
    check(
        "the four fields a ticket needs are all required",
        set(schema.get("required", []))
        >= {"addressed_to", "what_happened", "situation_summary", "next_steps"},
        f"required: {sorted(schema.get('required', []))}",
    )
    check(
        "every field tells the model what belongs in it",
        all(field.get("description") for field in properties.values()),
        "the description IS the instruction — there is nowhere else to say it",
    )

    with tempfile.TemporaryDirectory() as workspace:
        import os

        saved = os.environ.get("TICKETS_DIR")
        os.environ["TICKETS_DIR"] = workspace
        try:
            # Discovered again rather than reusing the cached handle. A tool
            # carries the environment its server will be spawned with, captured
            # when it was discovered — so the cached one would ignore
            # TICKETS_DIR and write a stray ticket into the real data/tickets.
            # It did exactly that the first time this check was written.
            fresh = asyncio.run(ticket_load_tools())
            writer = next(t for t in fresh if t.name == CREATE_TICKET)
            result = str(
                asyncio.run(
                    writer.ainvoke(
                        {
                            "addressed_to": "Housekeeping Supervisor",
                            "subject": "Round trip",
                            "what_happened": "A check called the tool.",
                            "situation_summary": "It should have written a file.",
                            "next_steps": ["Read the file back."],
                        }
                    )
                )
            )
            written = list(Path(workspace).glob("*.docx"))
            check("calling it over MCP writes a real file", len(written) == 1, result[:80])
            check(
                "it reports the path it wrote to",
                bool(written) and written[0].stem in result,
            )

            # Round trip: the node reads the destination back out of the tool's
            # real output. Checking the regex against a string written here
            # would only prove the regex matches itself — the wording that has
            # to keep matching belongs to the server, so it is the server's
            # output that gets parsed.
            raw = asyncio.run(
                writer.ainvoke(
                    {
                        "addressed_to": "Housekeeping Supervisor",
                        "subject": "Round trip parse",
                        "what_happened": "A check called the tool.",
                        "situation_summary": "The node must find the path in this.",
                        "next_steps": ["Parse it back."],
                    }
                )
            )
            parsed = fil._saved_path(
                [ToolMessage(raw, tool_call_id="1", name=CREATE_TICKET)]
            )
            check(
                "the filer can read the path back out of the tool's own result",
                bool(parsed) and Path(workspace).name in str(parsed),
                str(parsed),
            )
        finally:
            if saved is None:
                os.environ.pop("TICKETS_DIR", None)
            else:
                os.environ["TICKETS_DIR"] = saved

    # --- the node ---------------------------------------------------------
    check(
        "the filer is capped at one ticket per turn",
        fil.MAX_TICKETS == 1,
        "a second call would duplicate an incident already in the queue",
    )
    filed = [ToolMessage("done", tool_call_id="1", name=CREATE_TICKET)]
    check(
        "a filed ticket is counted from the tool's own result",
        fil._tickets_filed(filed) == 1
        and fil._tickets_filed([ToolMessage("x", tool_call_id="2", name="other")]) == 0,
    )
    check(
        "at the cap it is told to stop, and below it to file",
        "Do not call it again" in fil._note(fil.MAX_TICKETS)
        and "File the ticket now" in fil._note(0),
    )
    check(
        "the filer writes no References section",
        "do not write a `### References` section" in flat(fil.PROMPT),
        "a confirmation of an action cites nothing, and the verifier never sees it",
    )
    check(
        "the filer keeps the conversation going",
        "ask what else they need" in flat(fil.PROMPT),
        "filing is not a sign-off",
    )
    check(
        "the orchestrator is told never to file unasked",
        "only when the user has asked for a ticket" in flat(orch.PROMPT),
    )

    # --- when the ticket is offered ---------------------------------------
    #
    # Every case below was a real reply the agent gave. Offering after "Hello",
    # and after an answer that had just said it could not find a responsible
    # role, is what these guard against: the filer would have refused both, so
    # the offer promised something the rest of the workflow would not do.
    cli = load_cli()
    supervisor = "Housekeeping Supervisor (Housekeeping)"

    def reply(**fields) -> AgentState:
        return AgentState(
            **{
                "answer": f"Report it to the {supervisor}.",
                "delegations": ["researcher", RESPONDER],
                "verdict": "pass",
                "ticket_recipient": supervisor,
                **fields,
            }
        )

    check(
        "the ticket is offered when the answer establishes who to raise it with",
        cli.should_offer_ticket(reply()),
    )
    for label, fields, why in [
        ("after a greeting", {"answer": "Hello! How can I help?", "ticket_recipient": ""},
         "nothing to file and nobody to file it with"),
        ("when no responsible role was found", {"ticket_recipient": ""},
         "the answer said the documents do not name one"),
        ("straight after filing one", {"delegations": [RESPONDER, FILER]},
         "the filer asks its own follow-up"),
        ("when nothing was answered", {"answer": None}, ""),
        ("on an answer the verifier rejected", {"verdict": "fail"},
         "we have just warned that it is not grounded"),
        ("on an answer nothing graded", {"verdict": "ungraded"}, ""),
    ]:
        check(f"it is not offered {label}", not cli.should_offer_ticket(reply(**fields)), why)

    check(
        "the offer names who the ticket would go to",
        supervisor in cli.ticket_offer(supervisor)
        and "data/tickets" in cli.ticket_offer(supervisor),
        cli.ticket_offer(supervisor),
    )

    # --- who composes the reply, and who reads it -------------------------
    #
    # The workflow now builds the user-facing string once, at the node that ends
    # the run, and both front ends read it. These check that reading it and
    # composing it agree, and that a state carrying one is not silently
    # recomposed into a different one.
    composed = cli.reply_for(reply())
    check(
        "the composed reply is the answer followed by the offer",
        composed.startswith(f"Report it to the {supervisor}.")
        and cli.ticket_offer(supervisor) in composed,
    )
    check(
        "a front end reads the reply the graph built",
        cli.final_reply(reply(reply="what the user saw")) == "what the user saw",
        "the string shown and the string remembered have to be the same one",
    )
    check(
        "a run that ended without one is still answerable",
        cli.final_reply(reply()) == composed,
        "a front end printing nothing is worse than a recomputed string",
    )
    check(
        "the recipient is an observation, not a grading criterion",
        "does not affect your verdict" in flat(ver.Verdict.model_fields["ticket_recipient"].description)
        and "changes nothing about pass or fail" in flat(ver.PROMPT),
        "a verdict pulled by it would fail thin answers for having no recipient",
    )
    check(
        "the verifier clears the recipient when it rejects on citations",
        asyncio.run(
            ver.run(
                AgentState(
                    answer="Ask the Boss.\n\n### References\n- [1] made_up.pdf › page 9",
                    messages=[ToolMessage("x.pdf › page 1 › Role › Duties", tool_call_id="1")],
                    ticket_recipient="Boss",
                )
            )
        ).get("ticket_recipient")
        == "",
        "a partial update merges, so an earlier draft's recipient would survive",
    )


# ------------------------------------------------------------- observability


def check_observability() -> None:
    """Tracing must be invisible when off, and must not double-count when on.

    Both properties are decidable offline. The first matters because the agent
    has to run on a machine with no Langfuse; the second because LightRAG
    installs a global OpenAI patch whenever the keys are present, and the
    duplicate generations it produces would silently double every cost and token
    figure — measured, not assumed, see observability.py.
    """
    print("\nObservability (Langfuse)")
    import os

    from agent import observability as obs

    saved = {k: os.environ.get(k) for k in ("LANGFUSE_PUBLIC_KEY", "LANGFUSE_SECRET_KEY")}
    try:
        for key in saved:
            os.environ.pop(key, None)
        check("no keys means tracing is off", not obs.is_enabled())

        os.environ["LANGFUSE_PUBLIC_KEY"] = "pk-lf-check"
        check("one key alone is not enough", not obs.is_enabled())
        os.environ["LANGFUSE_SECRET_KEY"] = "sk-lf-check"
        check("both keys turn it on", obs.is_enabled())
        os.environ["LANGFUSE_TRACING_ENABLED"] = "false"
        check("the kill switch wins over the keys", not obs.is_enabled())
    finally:
        os.environ.pop("LANGFUSE_TRACING_ENABLED", None)
        for key, value in saved.items():
            os.environ.pop(key, None)
            if value is not None:
                os.environ[key] = value

    # The untraced path has to yield a usable object, not None — every caller
    # reads .callbacks off it unconditionally.
    #
    # Tracing is forced off here rather than assumed off. It was assumed once,
    # and the check passed until the day real keys landed in .env — at which
    # point it started failing on a machine where nothing was wrong. A check
    # that only holds in one configuration is testing the configuration.
    saved_client, saved_built = obs._client, obs._client_built
    os.environ["LANGFUSE_TRACING_ENABLED"] = "false"
    obs._client, obs._client_built = None, False
    try:
        with obs.trace_run("question") as trace:
            check("an untraced run still yields a handle", trace.callbacks == [])
            check("recording an untraced run is a no-op", trace.record(AgentState()) is None)
            check("an untraced run has no URL", trace.url is None)
    finally:
        os.environ.pop("LANGFUSE_TRACING_ENABLED", None)
        obs._client, obs._client_built = saved_client, saved_built

    # And the traced path, when this machine is set up for it. Skipped rather
    # than failed otherwise: not having Langfuse configured is a valid state.
    if obs.is_enabled():
        with obs.trace_run("question", tags=["check"]) as trace:
            check(
                "a configured run gets a callback handler",
                len(trace.callbacks) == 1,
                type(trace.callbacks[0]).__name__ if trace.callbacks else "none",
            )
        # Declaring the filter is not the same as it being applied: Langfuse
        # discards the argument when a client for this key already exists, and
        # says nothing. Asked here, after the client has actually been built.
        check(
            "the span filter is the one actually in force",
            obs.filter_in_force() is not False,
            "something built the Langfuse client before agent.observability did"
            if obs.filter_in_force() is False
            else "declared and applied",
        )
        deep = os.getenv("LANGFUSE_TRACE_OPENAI_SDK")
        if obs._truthy(deep):
            print(
                "      note: LANGFUSE_TRACE_OPENAI_SDK is on — LightRAG's internal "
                "calls are traced, and the agent's own LLM calls are recorded "
                "twice, so token and cost figures read about double."
            )
    else:
        print("      (traced path not checked — no Langfuse keys configured)")

    # The export filter, checked against stand-ins rather than live spans: the
    # only thing it reads is the span name, and building a real ReadableSpan
    # here would test OpenTelemetry rather than this decision.
    class _Span:
        def __init__(self, name):
            self.name = name
            self.attributes = {"langfuse.observation.type": "generation"}
            self.instrumentation_scope = type("S", (), {"name": "langfuse-sdk"})()
            self.parent = None

    os.environ.pop("LANGFUSE_TRACE_OPENAI_SDK", None)
    check(
        "the patched SDK's duplicate generations are dropped by default",
        not obs._should_export_span(_Span("OpenAI-generation")),
        "otherwise every agent LLM call is counted twice",
    )
    check(
        "the handler's own generations are kept",
        obs._should_export_span(_Span("ChatOpenAI")),
    )
    os.environ["LANGFUSE_TRACE_OPENAI_SDK"] = "true"
    check(
        "the duplicates can be opted back in",
        obs._should_export_span(_Span("OpenAI-generation")),
        "for debugging LightRAG, at the cost of doubled token counts",
    )
    os.environ.pop("LANGFUSE_TRACE_OPENAI_SDK", None)

    ok, message = obs.auth_check()
    check(
        "auth_check reports rather than raises",
        isinstance(ok, bool) and bool(message),
        message,
    )


def check_live(question: str) -> None:
    """One real call to the only implemented node."""
    print("\nLive orchestrator call (needs a working API key)")
    state = AgentState(messages=[HumanMessage(question)])
    try:
        result = asyncio.run(orch.run(state))
    except Exception as error:  # noqa: BLE001 - the point is to report it
        check("the orchestrator routed a real question", False, f"{type(error).__name__}: {error}")
        return
    check("the orchestrator routed a real question", result["next_node"] in WORKERS)
    print(f"      question: {question}")
    print(f"      routed to: {result['next_node']}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Check the agent's wiring.")
    parser.add_argument("--live", action="store_true", help="Also make one real model call.")
    parser.add_argument(
        "--question",
        default="Who has to approve a 700 dollar refund?",
        help="The question used for --live.",
    )
    args = parser.parse_args()

    check_structure()
    check_tool_boundaries()
    check_routers()
    check_orchestrator()
    check_retrieval_cascade()
    check_researcher()
    check_responder()
    check_verifier()
    check_declarations()
    check_tool_arguments()
    check_corpora()
    check_retrieval()
    check_full_lap()
    check_tickets()
    check_observability()
    if args.live:
        check_live(args.question)

    print(f"\n{len(PASSED)} passed, {len(FAILED)} failed")
    for label in FAILED:
        print(f"  FAILED: {label}")
    sys.exit(1 if FAILED else 0)


if __name__ == "__main__":
    main()
