"""The ticket document: a `Ticket` in, a saved `.docx` out.

No MCP and no model in here on purpose. The layout of the file is a rendering
decision, and keeping it separate from the protocol surface means it can be
called, read and checked without starting a server.

The template is fixed rather than composed by the caller. A generic
document-authoring tool — add_heading, add_paragraph, add_table, one call at a
time — would let the model lay out each ticket differently, and a ticket whose
shape depends on which model wrote it is not a form anyone can process. Here the
model supplies the *contents* of named fields and nothing else.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Project root: src/ticket_mcp/document.py -> src/ -> project/. Resolved from
# __file__ rather than the working directory, because this module runs inside a
# subprocess whose cwd is whatever the client that spawned it chose.
ROOT_DIR = Path(__file__).resolve().parents[2]

PRIORITIES = ("low", "normal", "high", "urgent")


def tickets_dir() -> Path:
    """Where tickets are written. `TICKETS_DIR` overrides, default data/tickets.

    Read on every call rather than captured at import, so a caller that sets the
    variable before spawning the server — or a test that points it at a temp
    directory — is actually obeyed.
    """
    value = os.getenv("TICKETS_DIR", "data/tickets")
    path = Path(value)
    return path if path.is_absolute() else ROOT_DIR / path


@dataclass(frozen=True)
class Ticket:
    """One ticket, as the fields a form would ask for.

    The four required fields are the four things a ticket has to carry to be
    actionable: who it is for, what happened, the situation around it, and what
    to do next. Everything else sharpens it and may legitimately be unknown —
    an empty optional is left out of the document rather than printed blank.
    """

    addressed_to: str
    subject: str
    what_happened: str
    situation_summary: str
    next_steps: tuple[str, ...] = ()
    raised_by: str = ""
    organisation: str = ""
    priority: str = "normal"
    references: tuple[str, ...] = ()
    raised_at: datetime = field(default_factory=datetime.now)

    def __post_init__(self) -> None:
        for name in ("addressed_to", "subject", "what_happened", "situation_summary"):
            if not str(getattr(self, name)).strip():
                raise ValueError(f"{name} is required and cannot be blank")
        if self.priority not in PRIORITIES:
            raise ValueError(f"priority must be one of {', '.join(PRIORITIES)}")

    @property
    def reference(self) -> str:
        """The human-facing ticket id, and the stem of its filename."""
        return f"TKT-{self.raised_at:%Y%m%d-%H%M%S}"


_UNSAFE = re.compile(r"[^A-Za-z0-9]+")


def _slug(text: str, limit: int = 48) -> str:
    """A filename-safe fragment of `text`.

    Windows rejects a longer list of characters than POSIX does and treats a few
    stems as devices, so everything outside [A-Za-z0-9] goes, rather than a
    denylist that would need to be right about all of them.
    """
    slug = _UNSAFE.sub("-", text).strip("-").lower()
    return slug[:limit].strip("-") or "ticket"


def _unique_path(directory: Path, stem: str) -> Path:
    """`stem.docx` in `directory`, suffixed if that name is taken.

    The id carries a whole-second timestamp, so two tickets filed in the same
    second would otherwise silently overwrite each other — and the one that
    disappears is the earlier report of the same incident.
    """
    candidate = directory / f"{stem}.docx"
    counter = 2
    while candidate.exists():
        candidate = directory / f"{stem}-{counter}.docx"
        counter += 1
    return candidate


def _field_table(document: Document, rows: list[tuple[str, str]]) -> None:
    """The header block: a two-column label/value table, blank rows dropped."""
    rows = [(label, value) for label, value in rows if str(value).strip()]
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        cells = table.rows[index].cells
        cells[0].text = ""
        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        cells[1].text = str(value)


def _section(document: Document, title: str, body: str) -> None:
    """One titled block of prose."""
    document.add_heading(title, level=2)
    document.add_paragraph(body.strip())


def write_ticket(ticket: Ticket, directory: Path | None = None) -> Path:
    """Render `ticket` to a .docx and return where it landed.

    The directory is created if it does not exist: the first ticket on a fresh
    checkout is the common case, and failing it would report a broken tool for
    what is really an empty data/ directory.
    """
    directory = directory or tickets_dir()
    directory.mkdir(parents=True, exist_ok=True)

    document = Document()

    document.add_heading("Incident Ticket", level=0)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(ticket.reference)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    _field_table(
        document,
        [
            ("Addressed to", ticket.addressed_to),
            ("Raised by", ticket.raised_by),
            ("Organisation", ticket.organisation),
            ("Priority", ticket.priority.capitalize()),
            ("Raised at", f"{ticket.raised_at:%Y-%m-%d %H:%M}"),
        ],
    )
    document.add_paragraph()

    _section(document, "Subject", ticket.subject)
    _section(document, "What happened", ticket.what_happened)
    _section(document, "Summary of the situation", ticket.situation_summary)

    if ticket.next_steps:
        document.add_heading("Next steps", level=2)
        for step in ticket.next_steps:
            document.add_paragraph(step.strip(), style="List Number")

    # The source labels ride along so the ticket can be checked against the
    # documents it was written from. A ticket that names a role but not the
    # section that made that role responsible is an assertion, not a record.
    if ticket.references:
        document.add_heading("Source references", level=2)
        for reference in ticket.references:
            document.add_paragraph(reference.strip(), style="List Bullet")

    footer = document.add_paragraph()
    note = footer.add_run(
        "Filed by the Role-Duty agent from the role and duty documents. "
        "Confirm against the source sections before acting on it."
    )
    note.italic = True
    note.font.size = Pt(8)
    note.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    path = _unique_path(directory, f"{ticket.reference}-{_slug(ticket.subject)}")
    document.save(path)
    return path
